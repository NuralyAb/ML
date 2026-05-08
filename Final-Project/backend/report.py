"""
Ministry-grade DOCX report generator.

Builds a Word document containing:
  * Title block with regional context and date
  * Executive summary (LLM-generated narrative — falls back to a deterministic
    text if OPENAI_API_KEY is missing or the call fails)
  * Forecast table for top-N diagnoses (recipe count by month)
  * Decisions / recommendations per disease tier
  * Methodology footer

The document is fully native Cyrillic — `python-docx` writes Open XML which
Word renders with the OS default Cyrillic-aware font.
"""
from __future__ import annotations
import io
import os
from dataclasses import dataclass
from datetime import datetime
from typing import Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))


@dataclass
class ForecastRow:
    icdid: str
    nozology: str
    history_avg_12m: float
    forecast: list[dict]   # [{year_month, predicted}]
    forecast_avg: float
    delta_pct: float
    decision: str          # surge / growth / stable / decline / drop


def _decision(delta: float) -> str:
    if delta < -15: return "drop"
    if delta < -5:  return "decline"
    if delta < 5:   return "stable"
    if delta < 15:  return "growth"
    return "surge"


def build_forecast_rows(panel_df: pd.DataFrame, forecaster, region: Optional[str], horizon: int, top_n: int) -> list[ForecastRow]:
    """Pick top diagnoses (by recipe count last 12 months) and forecast each."""
    last_m = panel_df["year_month"].max()
    cutoff = last_m - pd.DateOffset(months=12)
    recent = panel_df[panel_df["year_month"] > cutoff]
    if region:
        recent = recent[recent["region"] == region]
    top = (
        recent.groupby(["icdid", "nozology"], as_index=False)["recipe_count"].sum()
        .sort_values("recipe_count", ascending=False)
        .head(top_n)
    )

    rows: list[ForecastRow] = []
    for _, r in top.iterrows():
        icd = r["icdid"]; nozology = r["nozology"]

        if region:
            try:
                fc_df = forecaster.forecast(region, icd, horizon)
                forecast = fc_df.drop(columns=["nozology"]).to_dict(orient="records")
            except Exception:
                continue
            hist = forecaster.history_for(region, icd)
            base = float(hist.tail(12)["recipe_count"].mean())
        else:
            # Country-level: align by *calendar month*. Different regions stop
            # at different months in the source data; we forecast each region
            # forward enough to cover the same set of N months, then sum.
            sub = panel_df[panel_df["icdid"] == icd]
            country_hist = sub.groupby("year_month")["recipe_count"].sum().sort_index()
            base = float(country_hist.tail(12).mean())

            global_last = sub["year_month"].max()
            target_months = [global_last + pd.DateOffset(months=int(i)) for i in range(1, horizon + 1)]
            target_keys = [m.strftime("%Y-%m-%d") for m in target_months]

            sums = {k: 0.0 for k in target_keys}
            counts = {k: 0 for k in target_keys}
            for reg in sub["region"].unique():
                reg_last = sub[sub["region"] == reg]["year_month"].max()
                # How many steps do we need from this region to reach `global_last + horizon`?
                steps = int((target_months[-1].year - reg_last.year) * 12
                             + (target_months[-1].month - reg_last.month))
                if steps <= 0:
                    continue
                steps = min(steps, 18)  # safety cap on recursive drift
                try:
                    df_fc = forecaster.forecast(reg, icd, steps)
                except Exception:
                    continue
                for _, r in df_fc.iterrows():
                    key = pd.to_datetime(r["year_month"]).strftime("%Y-%m-%d")
                    if key in sums:
                        sums[key] += float(r["predicted"])
                        counts[key] += 1
            forecast = [
                {"year_month": k, "predicted": round(sums[k], 1)}
                for k in target_keys if counts[k] > 0
            ]

        if not forecast:
            continue
        f_avg = float(np.mean([f["predicted"] for f in forecast]))
        delta = (f_avg / base - 1) * 100 if base > 0 else 0.0
        rows.append(ForecastRow(
            icdid=icd, nozology=nozology, history_avg_12m=base,
            forecast=forecast, forecast_avg=f_avg, delta_pct=delta,
            decision=_decision(delta),
        ))
    return rows


# ---------------------------------------------------------------------------
# LLM narrative
# ---------------------------------------------------------------------------

DECISION_RU = {
    "surge":   "резкий рост",
    "growth":  "умеренный рост",
    "stable":  "стабильно",
    "decline": "снижение",
    "drop":    "резкий спад",
}


def llm_executive_summary(rows: list[ForecastRow], region: Optional[str], horizon: int) -> str:
    """Generate an executive summary via OpenAI; fall back to a deterministic
    sentence if the API key is missing or the call fails."""
    if not rows:
        return "Недостаточно данных для построения прогноза."

    grouped: dict[str, list[ForecastRow]] = {"surge": [], "growth": [], "stable": [], "decline": [], "drop": []}
    for r in rows:
        grouped[r.decision].append(r)

    bullets = []
    for r in rows:
        bullets.append(
            f"- {r.icdid} ({r.nozology}): база 12 мес. ≈ {round(r.history_avg_12m)}, "
            f"прогноз/мес ≈ {round(r.forecast_avg)}, Δ = {r.delta_pct:+.1f}%, "
            f"тренд = {DECISION_RU[r.decision]}"
        )
    prompt_data = "\n".join(bullets)

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return _fallback_summary(rows, region)

    try:
        # Lazy import — avoids hard dep when key is missing.
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        model = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

        scope = f"по региону «{region}»" if region else "по Республике Казахстан в целом"
        sys = (
            "Ты — аналитик Министерства здравоохранения РК. Готовишь сжатый раздел "
            "«Резюме для руководства» (executive summary) на русском языке. Стиль — "
            "деловой, без воды. Структура: 2–3 коротких абзаца. Цифры округляй до целых."
        )
        usr = (
            f"На основе следующих ML-прогнозов на {horizon} месяцев вперёд {scope} "
            f"составь резюме для руководства МЗ РК. Перечисли ключевые риски (рост спроса), "
            f"точки контроля (резкие спады), и две-три управленческие рекомендации "
            f"(закупки / кадры / аудит). Не упоминай конкретный ML-метод.\n\n"
            f"Прогнозы:\n{prompt_data}"
        )
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": sys}, {"role": "user", "content": usr}],
            temperature=0.3,
            max_tokens=700,
        )
        text = (resp.choices[0].message.content or "").strip()
        return text or _fallback_summary(rows, region)
    except Exception as e:  # noqa: BLE001
        print(f"[report] OpenAI call failed: {e}", flush=True)
        return _fallback_summary(rows, region)


def _fallback_summary(rows: list[ForecastRow], region: Optional[str]) -> str:
    surge = [r for r in rows if r.decision in ("surge", "growth")]
    drop = [r for r in rows if r.decision in ("drop", "decline")]
    scope = f"в регионе «{region}»" if region else "в Республике Казахстан"

    parts = [
        f"Прогнозный анализ объёмов выписки рецептов {scope} построен на "
        f"глобальной ML-модели (LightGBM), обученной на исторических данных "
        f"ЭРСБ МЗ РК с 2018 года. Ниже представлены ключевые выводы по "
        f"топ-{len(rows)} диагнозам по объёму выписки за последние 12 месяцев."
    ]
    if surge:
        names = ", ".join(f"{r.icdid} ({r.nozology})" for r in surge[:3])
        parts.append(
            "Ожидается рост спроса по следующим направлениям: " + names + ". "
            "Рекомендуется заблаговременно нарастить страховой запас препаратов "
            "и оценить пропускную способность поликлиник."
        )
    if drop:
        names = ", ".join(f"{r.icdid} ({r.nozology})" for r in drop[:3])
        parts.append(
            "Зафиксировано прогнозируемое снижение по: " + names + ". "
            "Рекомендуется аудит причин (возможные сбои регистрации) и "
            "перераспределение запасов между районами."
        )
    parts.append(
        "Подробная разбивка по диагнозам и месяцам приведена в таблице ниже."
    )
    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------

BRAND = RGBColor(0x1d, 0x4e, 0xd8)   # blue
INK   = RGBColor(0x0f, 0x17, 0x2a)
MUTED = RGBColor(0x64, 0x74, 0x8b)
SUCCESS = RGBColor(0x05, 0x96, 0x69)
WARNING = RGBColor(0xb4, 0x53, 0x09)
DANGER  = RGBColor(0xb1, 0x1d, 0x1d)


def _decision_color(d: str) -> RGBColor:
    return {
        "surge":   DANGER,
        "growth":  WARNING,
        "stable":  SUCCESS,
        "decline": MUTED,
        "drop":    BRAND,
    }.get(d, INK)


def _shade_cell(cell, hex_color: str):
    """Apply a background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def render_docx(rows: list[ForecastRow], summary: str, region: Optional[str], horizon: int, panel_df: pd.DataFrame) -> bytes:
    doc = Document()

    # Page setup: A4, narrow margins.
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Default font.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # Title.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("МИНИСТЕРСТВО ЗДРАВООХРАНЕНИЯ РЕСПУБЛИКИ КАЗАХСТАН")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    h = doc.add_paragraph()
    rh = h.add_run("Прогнозный отчёт по выписке рецептов")
    rh.font.bold = True
    rh.font.size = Pt(20)
    rh.font.color.rgb = INK

    sub = doc.add_paragraph()
    sr = sub.add_run(
        f"{('Регион: ' + region) if region else 'Республика Казахстан'} · "
        f"горизонт прогноза: {horizon} мес. · подготовлено: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = MUTED

    last_m = panel_df["year_month"].max()
    info = doc.add_paragraph()
    ir = info.add_run(
        f"Базовый период: исторические данные с {panel_df['year_month'].min().strftime('%Y-%m')} "
        f"по {last_m.strftime('%Y-%m')}; "
        f"источник: ЭРСБ МЗ РК (ISLO_MEDICALHISTORYOFCITIZENS)."
    )
    ir.font.size = Pt(9)
    ir.font.color.rgb = MUTED

    doc.add_paragraph()

    # Executive summary heading.
    p = doc.add_paragraph()
    r = p.add_run("Резюме для руководства")
    r.font.bold = True; r.font.size = Pt(14); r.font.color.rgb = BRAND

    for paragraph in summary.split("\n\n"):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # Forecast table.
    p = doc.add_paragraph()
    r = p.add_run("Прогноз по основным диагнозам")
    r.font.bold = True; r.font.size = Pt(14); r.font.color.rgb = BRAND

    if rows:
        forecast_months = [f["year_month"] for f in rows[0].forecast]
        cols = ["№", "МКБ", "Диагноз", "Среднее 12 мес."] + [
            datetime.strptime(m, "%Y-%m-%d").strftime("%b %Y") for m in forecast_months
        ] + ["Δ %", "Тренд"]
        table = doc.add_table(rows=1 + len(rows), cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"

        # Header.
        hdr = table.rows[0].cells
        for i, name in enumerate(cols):
            hdr[i].text = name
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
            _shade_cell(hdr[i], "1d4ed8")
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

        # Body.
        for i, fr in enumerate(rows, 1):
            cells = table.rows[i].cells
            cells[0].text = str(i)
            cells[1].text = fr.icdid
            cells[2].text = (fr.nozology or "")[:60]
            cells[3].text = f"{round(fr.history_avg_12m):,}".replace(",", " ")
            for j, m in enumerate(forecast_months):
                v = next((f["predicted"] for f in fr.forecast if f["year_month"] == m), 0)
                cells[4 + j].text = f"{round(float(v)):,}".replace(",", " ")
            cells[-2].text = f"{fr.delta_pct:+.1f}%"
            cells[-1].text = DECISION_RU[fr.decision]

            # Color the trend cell by tier.
            color_hex = {
                "surge": "fee2e2", "growth": "fef3c7",
                "stable": "d1fae5", "decline": "e0f2fe", "drop": "ede9fe",
            }[fr.decision]
            _shade_cell(cells[-1], color_hex)

            for c in cells:
                for paragraph in c.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()

    # Recommendations block.
    p = doc.add_paragraph()
    r = p.add_run("Рекомендации по управленческим действиям")
    r.font.bold = True; r.font.size = Pt(14); r.font.color.rgb = BRAND

    grouped: dict[str, list[ForecastRow]] = {k: [] for k in ("surge", "growth", "stable", "decline", "drop")}
    for fr in rows:
        grouped[fr.decision].append(fr)

    if grouped["surge"]:
        _add_block(doc, "Срочные действия (резкий рост)", DANGER, [
            f"{r.icdid} ({r.nozology}): {r.delta_pct:+.0f}% — "
            "нарастить запасы, привлечь резерв медперсонала, проинформировать центр закупок."
            for r in grouped["surge"][:5]
        ])
    if grouped["growth"]:
        _add_block(doc, "Плановые действия (умеренный рост)", WARNING, [
            f"{r.icdid} ({r.nozology}): {r.delta_pct:+.0f}% — "
            "увеличить страховой запас на 10–20%, проверить пропускную способность поликлиник."
            for r in grouped["growth"][:5]
        ])
    if grouped["drop"]:
        _add_block(doc, "Аудит (резкий спад)", BRAND, [
            f"{r.icdid} ({r.nozology}): {r.delta_pct:+.0f}% — "
            "проверить корректность регистрации, перенаправить ресурсы."
            for r in grouped["drop"][:5]
        ])
    if grouped["decline"]:
        _add_block(doc, "Оптимизация (умеренное снижение)", MUTED, [
            f"{r.icdid} ({r.nozology}): {r.delta_pct:+.0f}% — "
            "сократить новые закупки, перераспределить запасы между районами."
            for r in grouped["decline"][:5]
        ])

    # Methodology footer.
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Методология")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = MUTED

    methodology = (
        "Прогноз построен глобальной градиентно-бустинговой моделью (LightGBM), "
        "обученной на агрегированных по месяцам данных ЭРСБ МЗ РК (ISLO_MEDICALHISTORYOFCITIZENS) "
        "за период 2018–2025 гг. (>100 млн рецептов, 18 регионов, ≈2 300 кодов МКБ-10). "
        "Признаки модели: лаги 1–12 месяцев, скользящие средние и стандартные отклонения, "
        "сезонные синусы/косинусы, доля региона в общем потоке. "
        "Метрики на отложенной выборке (последние 6 месяцев каждой серии): "
        "MAE 18.7, RMSE 278.8, R² 0.94 (LightGBM выигрывает у наивных бейзлайнов "
        "на 36% по MAE). Многошаговый прогноз — рекурсивный."
    )
    p = doc.add_paragraph(methodology)
    for run in p.runs:
        run.font.size = Pt(9); run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    r = p.add_run(
        "Резюме для руководства сгенерировано с использованием большой языковой модели "
        "на основе численных результатов прогноза. Решения по закупкам, кадрам и аудиту "
        "должны приниматься уполномоченными лицами с учётом полной картины состояния системы."
    )
    for run in p.runs:
        run.font.size = Pt(8); run.font.italic = True; run.font.color.rgb = MUTED

    # Serialize.
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_block(doc, title: str, color: RGBColor, lines: list[str]):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = color
    for line in lines:
        bp = doc.add_paragraph(line, style="List Bullet")
        for run in bp.runs:
            run.font.size = Pt(10)
